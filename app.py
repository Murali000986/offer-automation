import os
import csv
import json
import io
import zipfile
import datetime
import pandas as pd
import traceback # For detailed error logging
import shutil # For copying files to template library
from flask import (
    Flask, render_template, request, send_from_directory,
    flash, redirect, url_for, send_file
)
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Pt # Optional: Could be used for advanced formatting if needed

# --- DOCX2PDF Import and Availability Check ---
try:
    from docx2pdf import convert as docx_to_pdf_convert
    DOCX2PDF_AVAILABLE = True
    print("✅ docx2pdf library found and imported successfully.")
except ImportError:
    print("⚠️ WARNING: `docx2pdf` library not found. PDF generation will be disabled.")
    DOCX2PDF_AVAILABLE = False
    docx_to_pdf_convert = None # Define as None for graceful checks later
except Exception as import_err:
    print(f"⚠️ WARNING: Error importing `docx2pdf`: {import_err}. PDF generation might fail.")
    DOCX2PDF_AVAILABLE = False
    docx_to_pdf_convert = None

# === Flask App Setup ===
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", os.urandom(24))

# === Configuration ===
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'temp_uploads') # For temporary storage before processing
GENERATED_FOLDER = os.path.join(BASE_DIR, 'generated_files') # Where final DOCX/PDFs are saved
TEMPLATES_FOLDER = os.path.join(BASE_DIR, 'user_templates') # For persistent user templates
ALLOWED_EXTENSIONS_DOCX = {'docx'}
ALLOWED_EXTENSIONS_DATA = {'csv', 'json'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GENERATED_FOLDER'] = GENERATED_FOLDER
app.config['TEMPLATES_FOLDER'] = TEMPLATES_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 32 * 1024 * 1024 # 32 MB Upload limit

# Ensure necessary folders exist with write permissions
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)
os.makedirs(TEMPLATES_FOLDER, exist_ok=True)

# === Helper Functions ===

def allowed_file(filename, allowed_extensions):
    """Checks if a filename has an allowed extension (case-insensitive)."""
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

def replace_text_in_runs(paragraphs, data):
    """Replaces placeholders {key} with value from data dict in paragraph runs."""
    replacements = 0
    if not data: # Skip if data is empty
        return 0
    # Ensure all placeholders are strings (even if they look like {numbers})
    placeholders_map = {str(k): str(v) if v is not None else '' for k, v in data.items()}
    placeholders = list(placeholders_map.keys()) # Get keys once

    for paragraph in paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        if not any(ph in full_text for ph in placeholders):
            continue
        for run in paragraph.runs:
            if not any(ph_part in run.text for ph in placeholders for ph_part in ['{','}']):
                 continue
            original_run_text = run.text
            modified_run_text = original_run_text
            for placeholder, value in placeholders_map.items():
                 modified_run_text = modified_run_text.replace(placeholder, value)
            if original_run_text != modified_run_text:
                run.text = modified_run_text
                replacements += 1
    return replacements

def safe_cleanup(filepath):
    """Attempts to remove a file, logging errors but not crashing."""
    if filepath and os.path.exists(filepath):
        try:
            os.remove(filepath)
            print(f"DEBUG: Cleaned up temporary file: {filepath}")
        except OSError as e:
            print(f"Warning: Could not remove temporary file {filepath}: {e}")

def list_user_templates():
    """Lists .docx templates from the TEMPLATES_FOLDER."""
    if not os.path.exists(app.config['TEMPLATES_FOLDER']):
        return []
    templates = [f for f in os.listdir(app.config['TEMPLATES_FOLDER']) if f.lower().endswith('.docx')]
    return sorted(templates)

# === Core Generation Logic Functions ===

def generate_document_core(template_path, data, filename_prefix, filename_suffix):
    """
    Core logic to generate a DOCX and optionally PDF from a template.
    """
    print(f"DEBUG: Generating '{filename_prefix}' for suffix: {filename_suffix}")
    docx_filename = f"{filename_prefix}_{filename_suffix}.docx"
    docx_save_path = os.path.join(app.config['GENERATED_FOLDER'], docx_filename)
    pdf_filename = f"{filename_prefix}_{filename_suffix}.pdf"
    pdf_save_path = os.path.join(app.config['GENERATED_FOLDER'], pdf_filename)
    pdf_final_path = None
    error_message = None

    try:
        if not os.path.exists(template_path):
            return False, None, None, f"Template file not found: '{template_path}'"
        if not data or not isinstance(data, dict):
            return False, None, None, "Invalid or missing data for document generation."

        print(f"DEBUG: Loading template: {template_path}")
        doc = Document(template_path)
        total_replacements = 0
        print(f"DEBUG: Replacing placeholders in {filename_prefix}...")
        total_replacements += replace_text_in_runs(doc.paragraphs, data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: total_replacements += replace_text_in_runs(cell.paragraphs, data)
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header is not None: total_replacements += replace_text_in_runs(header.paragraphs, data)
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer is not None: total_replacements += replace_text_in_runs(footer.paragraphs, data)

        if total_replacements == 0:
            print(f"⚠️ Warning: No placeholders were replaced for '{filename_suffix}' ({filename_prefix}). Check template & data keys.")

        print(f"DEBUG: Saving DOCX to: {docx_save_path}")
        doc.save(docx_save_path)
        if not os.path.exists(docx_save_path):
            raise OSError(f"Failed to save DOCX file to '{docx_save_path}'")
        print(f"✅ Generated DOCX: {docx_filename}")

        pdf_specific_error = None
        if DOCX2PDF_AVAILABLE and docx_to_pdf_convert:
            print(f"DEBUG: Attempting PDF conversion: {docx_filename} -> {pdf_filename}")
            try:
                docx_to_pdf_convert(docx_save_path, pdf_save_path)
                if os.path.exists(pdf_save_path):
                    pdf_final_path = pdf_save_path
                    print(f"✅ Generated PDF: {pdf_filename}")
                else:
                    pdf_specific_error = "PDF Conversion process completed, but the output file was not found."
                    print(f"❌ {pdf_specific_error}")
            except Exception as pdf_err:
                 pdf_specific_error = f"PDF conversion failed: {str(pdf_err)}"
                 print(f"❌ Error during PDF conversion for '{filename_suffix}': {pdf_err}")
                 flash(f"Warning: PDF generation failed for {filename_suffix}. DOCX created. Error: {str(pdf_err)}", "warning")
        else:
             if DOCX2PDF_AVAILABLE and not docx_to_pdf_convert:
                 print("ERROR: docx2pdf marked available but convert function missing.")
                 pdf_specific_error = "PDF conversion function unavailable despite library import."
             else:
                print(f"DEBUG: PDF conversion skipped for {filename_suffix} (library/function not available).")

        return True, docx_save_path, pdf_final_path, pdf_specific_error

    except Exception as e:
        error_message = f"Failed to generate {filename_prefix} for '{filename_suffix}': {str(e)}"
        print(f"❌ {error_message}")
        print(traceback.format_exc())
        safe_cleanup(docx_save_path) # Cleanup partially generated docx
        safe_cleanup(pdf_save_path) # Cleanup partially generated pdf
        return False, None, None, error_message

# Wrapper functions
def generate_offer_letter_web(template_path, data, filename_suffix):
     return generate_document_core(template_path, data, "Offer_Letter", filename_suffix)

def generate_relieving_letter_web(template_path, data, filename_suffix):
     return generate_document_core(template_path, data, "Relieving_Letter", filename_suffix)

# === Flask Routes ===

@app.context_processor
def inject_global_vars():
    return {
        'current_year': datetime.datetime.now().year,
        'user_templates': list_user_templates() # Make templates available to all renders
    }

@app.route('/')
def index():
    # Check if 'active_tab' is passed from a redirect and preserve it
    active_tab = request.args.get('active_tab', None)
    # Check if specific download info is passed and preserve it
    letter_type = request.args.get('letter_type', None)
    docx_download = request.args.get('docx_download', None)
    pdf_download = request.args.get('pdf_download', None)
    recipient_name = request.args.get('recipient_name', None)
    converter_pdf_download = request.args.get('converter_pdf_download', None)
    original_docx_filename = request.args.get('original_docx_filename', None)

    return render_template('index.html',
                           active_tab=active_tab,
                           letter_type=letter_type,
                           docx_download=docx_download,
                           pdf_download=pdf_download,
                           recipient_name=recipient_name,
                           converter_pdf_download=converter_pdf_download,
                           original_docx_filename=original_docx_filename)


# --- Template Management Routes ---
@app.route('/upload_user_template', methods=['POST'])
def handle_upload_user_template():
    active_tab_anchor = 'manage-templates-content'
    if 'user_template_file' not in request.files:
        flash('No template file provided for library upload.', 'danger')
        return redirect(url_for('index', active_tab=active_tab_anchor))
    
    file = request.files['user_template_file']
    if file.filename == '':
        flash('No file selected for library upload.', 'danger')
        return redirect(url_for('index', active_tab=active_tab_anchor))

    if file and allowed_file(file.filename, ALLOWED_EXTENSIONS_DOCX):
        filename = secure_filename(file.filename)
        save_path = os.path.join(app.config['TEMPLATES_FOLDER'], filename)
        
        if os.path.exists(save_path) and not request.form.get('overwrite_template'):
            flash(f"Template '{filename}' already exists. Check 'Overwrite' to replace it.", 'warning')
            return redirect(url_for('index', active_tab=active_tab_anchor))
            
        try:
            file.save(save_path)
            flash(f"Template '{filename}' uploaded successfully to library.", 'success')
        except Exception as e:
            flash(f"Error uploading template '{filename}': {e}", 'danger')
    else:
        flash('Invalid file type. Only .docx templates are allowed.', 'danger')
    return redirect(url_for('index', active_tab=active_tab_anchor))


@app.route('/delete_user_template/<path:filename>', methods=['POST'])
def handle_delete_user_template(filename):
    active_tab_anchor = 'manage-templates-content'
    safe_filename = secure_filename(filename)
    if not safe_filename:
        flash("Invalid template filename for deletion.", "danger")
        return redirect(url_for('index', active_tab=active_tab_anchor))

    template_path = os.path.join(app.config['TEMPLATES_FOLDER'], safe_filename)
    if os.path.exists(template_path):
        try:
            os.remove(template_path)
            flash(f"Template '{safe_filename}' deleted successfully.", 'success')
        except Exception as e:
            flash(f"Error deleting template '{safe_filename}': {e}", 'danger')
    else:
        flash(f"Template '{safe_filename}' not found for deletion.", 'warning')
    return redirect(url_for('index', active_tab=active_tab_anchor))


# --- Helper for processing template choice ---
def process_template_input(letter_type_for_log="document"):
    """
    Processes template input from form (existing library template or new upload).
    Returns (actual_template_path_to_use, temp_uploaded_path_for_cleanup_if_any)
    Returns (None, None) on error and flashes messages.
    """
    template_choice = request.form.get('template_choice')
    selected_template_filename = request.form.get('selected_template_filename')
    save_to_library = request.form.get('save_uploaded_template_to_library') == 'on'
    
    actual_template_path = None
    temp_uploaded_path = None

    if template_choice == 'existing_template':
        if not selected_template_filename:
            flash('Please select an existing template from the library.', 'danger')
            return None, None
        
        safe_selected_filename = secure_filename(os.path.basename(selected_template_filename))
        if not safe_selected_filename: 
            flash('Invalid selected template name.', 'danger')
            return None, None

        actual_template_path = os.path.join(app.config['TEMPLATES_FOLDER'], safe_selected_filename)
        if not os.path.exists(actual_template_path):
            flash(f'Selected library template "{safe_selected_filename}" not found.', 'danger')
            return None, None
        print(f"DEBUG: Using existing template: {actual_template_path}")

    elif template_choice == 'upload_new_template':
        if 'template_file' not in request.files or not request.files['template_file'].filename:
            flash(f'No template file provided for new upload ({letter_type_for_log}).', 'danger')
            return None, None
        
        template_file = request.files['template_file']
        if not allowed_file(template_file.filename, ALLOWED_EXTENSIONS_DOCX):
            flash(f'Invalid template file type for {letter_type_for_log} (.docx required).', 'danger')
            return None, None

        original_filename = secure_filename(template_file.filename)
        timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')
        
        temp_filename_for_processing = f"runtime_upload_{letter_type_for_log}_{timestamp}_{original_filename}"
        temp_uploaded_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename_for_processing)
        
        try:
            template_file.save(temp_uploaded_path)
            actual_template_path = temp_uploaded_path 
            print(f"DEBUG: Saved temporary uploaded template: {temp_uploaded_path}")
        except Exception as e:
            flash(f"Error saving temporary template: {e}", "danger")
            safe_cleanup(temp_uploaded_path)
            return None, None

        if save_to_library:
            library_path = os.path.join(app.config['TEMPLATES_FOLDER'], original_filename)
            try:
                if os.path.exists(library_path) and not request.form.get('overwrite_if_exists_in_library_from_form'): 
                     flash(f"Template '{original_filename}' already exists in library. Not overwritten from this form. Use 'Manage Templates' to explicitly overwrite.", 'warning')
                else:
                    shutil.copy(temp_uploaded_path, library_path)
                    flash(f"Template '{original_filename}' saved to library.", "info")
                    app.jinja_env.globals.update(user_templates=list_user_templates())
            except Exception as e:
                flash(f"Could not save uploaded template '{original_filename}' to library: {e}", "warning")
    else:
        flash('Invalid template choice. Please select or upload a template.', 'danger')
        return None, None
        
    return actual_template_path, temp_uploaded_path


# --- Single Generation Routes ---
@app.route('/generate_single_offer', methods=['POST'])
def handle_generate_single_offer():
    letter_type = 'offer'
    uploaded_temp_template_path_for_cleanup = None
    recipient_name_field = 'candidate name' # This is the HTML form field name
    
    # *** CORRECTED KEYS TO MATCH HTML UNDERSCORE CONVENTION WHERE APPLICABLE ***
    form_to_placeholder = {
        "send_date": "{send date}", "candidate name": "{candidate name}", "designation": "{designation}",
        "fdesignation": "{fdesignation}", "email": "{email}", 
        "mobile_number": "{mobile number}",  # Expects HTML name="mobile_number"
        "dear_name": "{dear name}",          # Expects HTML name="dear_name"
        "joining_date": "{joining date}", "hr_name": "{hr name}", "lpa": "{lpa}"
    }
    # The field used for filename suffix and display messages
    # This assumes 'candidate name' (with space) is correctly mapped if `recipient_name_field` is 'candidate name'
    recipient_name_from_form = request.form.get(recipient_name_field, f'Unknown_{letter_type}_Recipient').strip()
    success_flag = False 
    active_tab_anchor = f'{letter_type}-single-content'
    
    try:
        actual_template_to_use, uploaded_temp_template_path_for_cleanup = process_template_input(letter_type)
        if not actual_template_to_use:
            return redirect(url_for('index', active_tab=active_tab_anchor))

        form_data = {}
        missing_fields = []
        for field_key, placeholder in form_to_placeholder.items():
            # `field_key` is what we use with request.form.get()
            value = request.form.get(field_key, '').strip() 
            is_required = field_key != "fdesignation" 
            if not value and is_required:
                missing_fields.append(f"'{field_key.replace('_', ' ')}'") # Make field name more readable for flash
            form_data[placeholder] = value

        if missing_fields:
             flash(f"Missing required {letter_type} details: {', '.join(missing_fields)}.", 'danger')
             raise ValueError("Missing fields")

        name_suffix = recipient_name_from_form.replace(" ", "_").replace("/", "-").replace("\\", "-")
        output_timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename_suffix = secure_filename(f"single_{letter_type}_{name_suffix}_{output_timestamp}")

        print(f"--- Generating Single {letter_type.capitalize()}: {recipient_name_from_form} ---")
        gen_success, docx_path, pdf_path, error_msg = generate_offer_letter_web(
             actual_template_to_use, form_data, filename_suffix
        )

        if gen_success:
            success_flag = True
            flash(f"Successfully generated {letter_type} document(s) for {recipient_name_from_form}.", 'success')
            final_docx = os.path.basename(docx_path) if docx_path and os.path.exists(docx_path) else None
            final_pdf = os.path.basename(pdf_path) if pdf_path and os.path.exists(pdf_path) else None
            if not final_docx:
                 flash("Internal Error: Generated DOCX file not found.", "danger")
                 return redirect(url_for('index', active_tab=active_tab_anchor))
            # Pass results back to index page via query parameters for display
            return redirect(url_for('index', 
                                   letter_type=letter_type,
                                   docx_download=final_docx,
                                   pdf_download=final_pdf,
                                   recipient_name=recipient_name_from_form,
                                   active_tab=active_tab_anchor))
        else:
            flash(f"Error generating {letter_type} letter for {recipient_name_from_form}: {error_msg}", 'danger')
            return redirect(url_for('index', active_tab=active_tab_anchor))

    except ValueError: 
        pass 
    except Exception as route_err:
        print(f"CRITICAL Route Error in /generate_single_{letter_type}: {route_err}\n{traceback.format_exc()}")
        flash(f"Unexpected server error during single {letter_type} generation.", "danger")
    finally:
        safe_cleanup(uploaded_temp_template_path_for_cleanup)
        if not success_flag: 
             return redirect(url_for('index', active_tab=active_tab_anchor))


@app.route('/generate_single_relieving', methods=['POST'])
def handle_generate_single_relieving():
    letter_type = 'relieving'
    uploaded_temp_template_path_for_cleanup = None
    recipient_name_field = 'name' # HTML form field name
    form_to_placeholder = { # Keys are HTML form field names
        "send_date": "{send date}", "name": "{name}", "role": "{role}",
        "working_date": "{working date}", "accepted_date": "{accepted date}",
        "relieved_date": "{relieved date}", "hr_name": "{hr name}"
    }
    recipient_name_from_form = request.form.get(recipient_name_field, f'Unknown_{letter_type}_Recipient').strip()
    success_flag = False
    active_tab_anchor = f'{letter_type}-single-content'
    try:
        actual_template_to_use, uploaded_temp_template_path_for_cleanup = process_template_input(letter_type)
        if not actual_template_to_use:
            return redirect(url_for('index', active_tab=active_tab_anchor))

        form_data = {}
        missing_fields = []
        for field_key, placeholder in form_to_placeholder.items():
             value = request.form.get(field_key, '').strip()
             if not value: 
                  missing_fields.append(f"'{field_key.replace('_', ' ')}'")
             form_data[placeholder] = value
        if missing_fields:
             flash(f"Missing required {letter_type} details: {', '.join(missing_fields)}.", 'danger'); raise ValueError("Missing fields")

        name_suffix = recipient_name_from_form.replace(" ", "_").replace("/", "-").replace("\\", "-")
        output_timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
        filename_suffix = secure_filename(f"single_{letter_type}_{name_suffix}_{output_timestamp}")

        print(f"--- Generating Single {letter_type.capitalize()}: {recipient_name_from_form} ---")
        gen_success, docx_path, pdf_path, error_msg = generate_relieving_letter_web(
             actual_template_to_use, form_data, filename_suffix
        )

        if gen_success:
            success_flag = True
            flash(f"Successfully generated {letter_type} document(s) for {recipient_name_from_form}.", 'success')
            final_docx = os.path.basename(docx_path) if docx_path and os.path.exists(docx_path) else None
            final_pdf = os.path.basename(pdf_path) if pdf_path and os.path.exists(pdf_path) else None
            if not final_docx: flash("Internal Error: Generated DOCX file not found.", "danger"); return redirect(url_for('index', active_tab=active_tab_anchor))
            return redirect(url_for('index', 
                                   letter_type=letter_type, 
                                   docx_download=final_docx, 
                                   pdf_download=final_pdf, 
                                   recipient_name=recipient_name_from_form, 
                                   active_tab=active_tab_anchor))
        else: 
            flash(f"Error generating {letter_type} letter for {recipient_name_from_form}: {error_msg}", 'danger')
            return redirect(url_for('index', active_tab=active_tab_anchor))

    except ValueError:
        pass
    except Exception as route_err:
         print(f"CRITICAL Route Error in /generate_single_{letter_type}: {route_err}\n{traceback.format_exc()}"); flash(f"Unexpected server error during single {letter_type} generation.", "danger")
    finally:
        safe_cleanup(uploaded_temp_template_path_for_cleanup)
        if not success_flag:
             return redirect(url_for('index', active_tab=active_tab_anchor))


# --- Bulk Generation Routes ---

@app.route('/generate_bulk_manual_offer', methods=['POST'])
def handle_generate_bulk_manual_offer():
    letter_type = 'offer'
    uploaded_temp_template_path_for_cleanup = None
    zip_buffer = None 
    output_format = request.form.get('output_format', 'both')
    # Keys are HTML form field names (generated by JS with underscores)
    form_to_placeholder = {
        "candidate_name": "{candidate name}", "send_date": "{send date}", "designation": "{designation}",
        "fdesignation": "{fdesignation}", "email": "{email}", "mobile_number": "{mobile number}",
        "dear_name": "{dear name}", "joining_date": "{joining date}", "hr_name": "{hr name}", "lpa": "{lpa}"
    }
    required_fields = ['candidate_name', 'send_date', 'designation', 'email', 'mobile_number', 'dear_name', 'joining_date', 'hr_name', 'lpa'] # These are the form field name keys
    name_field_key = 'candidate_name' # This is a form field name key
    active_tab_anchor = 'offer-manual-content'

    try:
        actual_template_to_use, uploaded_temp_template_path_for_cleanup = process_template_input(f"manual_bulk_{letter_type}")
        if not actual_template_to_use:
            return redirect(url_for('index', active_tab=active_tab_anchor))

        form_lists = { name_key: request.form.getlist(name_key) for name_key in form_to_placeholder.keys() }
        num_entries = len(form_lists.get(name_field_key, [])) 

        if num_entries == 0: flash("No entry details submitted for manual bulk offer.", "warning"); raise ValueError("No entries")

        for field_name_check in form_to_placeholder.keys(): 
            if len(form_lists.get(field_name_check,[])) != num_entries:
                 flash(f"Data mismatch: Inconsistent number of entries for field '{field_name_check.replace('_', ' ')}'. Expected {num_entries}, got {len(form_lists.get(field_name_check,[]))}.", "danger")
                 raise ValueError("List length mismatch")
        print(f"DEBUG: Processing {num_entries} manually entered {letter_type} letters.")

        generated_files_paths = []; generation_errors = []; pdf_failures = 0; successful_records_count = 0;
        if output_format in ['pdf', 'both'] and not DOCX2PDF_AVAILABLE: flash("PDF conversion unavailable.", "warning")

        for i in range(num_entries):
            current_data_for_template = {} 
            current_missing_fields_keys = []
            recipient_name = form_lists[name_field_key][i].strip() 

            for field_key, placeholder in form_to_placeholder.items():
                value = form_lists[field_key][i].strip()
                current_data_for_template[placeholder] = value
                # fdesignation is optional (not in required_fields)
                is_required = field_key in required_fields and field_key != "fdesignation"
                if is_required and not value:
                    current_missing_fields_keys.append(f"'{field_key.replace('_', ' ')}'")
            
            # Special check for fdesignation if it IS in required_fields (it's not here, but good practice)
            if "fdesignation" in required_fields and not form_lists["fdesignation"][i].strip():
                current_missing_fields_keys.append("'fdesignation'")


            if not recipient_name:
                 err_msg = f"Manual {letter_type.capitalize()} Entry {i+1}: Skipping (Missing '{name_field_key.replace('_', ' ')}')"; print(f"Warning:{err_msg}"); generation_errors.append(err_msg); continue
            if current_missing_fields_keys:
                 err_msg = f"Manual {letter_type.capitalize()} Entry {i+1} ('{recipient_name}'): Missing fields: {', '.join(current_missing_fields_keys)}"; print(f"Warning:{err_msg}"); generation_errors.append(err_msg); continue

            name_suffix = recipient_name.replace(" ", "_").replace("/", "-").replace("\\", "-"); out_ts = datetime.datetime.now().strftime('%H%M%S%f')
            filename_suffix = secure_filename(f"manual_{letter_type}_{i+1}_{name_suffix}_{out_ts}")
            print(f"--- Processing Manual {letter_type.capitalize()} ({i+1}/{num_entries}): {recipient_name} ---")
            gen_success, docx_path, pdf_path, pdf_err_msg = generate_offer_letter_web(actual_template_to_use, current_data_for_template, filename_suffix)

            added_files_for_record = False
            if gen_success:
                 if output_format in ['docx', 'both'] and docx_path and os.path.exists(docx_path): generated_files_paths.append(docx_path); added_files_for_record = True
                 if output_format in ['pdf', 'both']:
                     if pdf_path and os.path.exists(pdf_path):
                         generated_files_paths.append(pdf_path); added_files_for_record = True
                     else:
                         pdf_failures += 1 
                         if pdf_err_msg: print(f"DEBUG: PDF failure reason for '{recipient_name}': {pdf_err_msg}")
                 if added_files_for_record:
                     successful_records_count += 1
                 else:
                     err = f"'{recipient_name}' (Manual {letter_type.capitalize()}): DOCX generated but no requested output format found/saved."
                     generation_errors.append(err); print(f"Warning: {err}") 
            else:
                fail_reason = pdf_err_msg or 'Generation core function returned False.'
                generation_errors.append(f"'{recipient_name}' (Manual {letter_type.capitalize()}): Generation failed ({fail_reason})")

        print(f"DEBUG: Manual bulk {letter_type} finished. Success: {successful_records_count}/{num_entries}")
        if successful_records_count == 0:
             flash(f"No {letter_type} letters generated successfully from manual entries.", "danger");
             if generation_errors: flash("Errors: " + "; ".join(generation_errors), "warning");
             raise ValueError("No successful generation")

        zip_buffer = io.BytesIO(); print(f"DEBUG: Creating Manual {letter_type.capitalize()} ZIP ({output_format})...");
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
             for file_path in generated_files_paths:
                 if os.path.exists(file_path): zipf.write(file_path, os.path.basename(file_path))
                 else: print(f"Warning: File path listed for zipping but not found: {file_path}") 

        zip_buffer.seek(0); success_msg = f"Generated documents for {successful_records_count}/{num_entries} manual {letter_type} entries (Format: {output_format.upper()})."; flash(success_msg, 'success')
        if generation_errors: flash("Issues: " + "; ".join(generation_errors), 'warning')
        if pdf_failures > 0 and output_format != 'docx': flash(f"Note: PDF failed/skipped for {pdf_failures} entries.", 'info')
        zip_filename_ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S"); zip_filename = f"Generated_Manual_Offers_{output_format.upper()}_{zip_filename_ts}.zip"
        # Do not redirect here, send_file will be the response
        return send_file( zip_buffer, mimetype='application/zip', as_attachment=True, download_name=zip_filename )

    except ValueError: 
        pass 
    except Exception as route_err:
        print(f"CRITICAL Error /generate_bulk_manual_{letter_type}: {route_err}\n{traceback.format_exc()}"); flash("Unexpected server error during manual bulk generation.", "danger")
    finally:
        safe_cleanup(uploaded_temp_template_path_for_cleanup)
        if 'generated_files_paths' in locals() and zip_buffer: # Only cleanup if zip was created
             for p_cleanup in generated_files_paths: safe_cleanup(p_cleanup)
        
        if 'zip_buffer' not in locals() or not zip_buffer: # Redirect if zip wasn't created/sent
             return redirect(url_for('index', active_tab=active_tab_anchor))


@app.route('/generate_bulk_manual_relieving', methods=['POST'])
def handle_generate_bulk_manual_relieving():
    letter_type = 'relieving'
    uploaded_temp_template_path_for_cleanup = None
    zip_buffer = None 
    output_format = request.form.get('output_format', 'both')
    # Keys are HTML form field names (generated by JS with underscores if applicable)
    form_to_placeholder = {
        "name": "{name}", "send_date": "{send date}", "role": "{role}",
        "working_date": "{working date}", "accepted_date": "{accepted date}",
        "relieved_date": "{relieved date}", "hr_name": "{hr name}"
    }
    required_fields_keys = list(form_to_placeholder.keys()) # All are required for relieving
    name_field_key = 'name' # This is a form field name key
    active_tab_anchor = 'relieving-manual-content'

    try:
        actual_template_to_use, uploaded_temp_template_path_for_cleanup = process_template_input(f"manual_bulk_{letter_type}")
        if not actual_template_to_use:
            return redirect(url_for('index', active_tab=active_tab_anchor))

        form_lists = { name_key: request.form.getlist(name_key) for name_key in form_to_placeholder.keys() }
        num_entries = len(form_lists.get(name_field_key, [])) 

        if num_entries == 0: flash("No entry details submitted for manual bulk relieving.", "warning"); raise ValueError("No entries")

        for field_name_check in form_to_placeholder.keys():
            if len(form_lists.get(field_name_check,[])) != num_entries:
                 flash(f"Data mismatch: Inconsistent number of entries for field '{field_name_check.replace('_', ' ')}'. Expected {num_entries}, got {len(form_lists[field_name_check])}.", "danger")
                 raise ValueError("List length mismatch")

        print(f"DEBUG: Processing {num_entries} manually entered {letter_type} letters.")

        generated_files_paths = []; generation_errors = []; pdf_failures = 0; successful_records_count = 0;
        if output_format in ['pdf', 'both'] and not DOCX2PDF_AVAILABLE: flash("PDF conversion unavailable.", "warning")

        for i in range(num_entries):
            current_data_for_template = {} 
            current_missing_fields_keys = []
            recipient_name = form_lists[name_field_key][i].strip() 

            for field_key, placeholder in form_to_placeholder.items():
                value = form_lists[field_key][i].strip()
                current_data_for_template[placeholder] = value
                if field_key in required_fields_keys and not value:
                    current_missing_fields_keys.append(f"'{field_key.replace('_', ' ')}'")

            if not recipient_name: 
                 err_msg = f"Manual {letter_type.capitalize()} Entry {i+1}: Skipping (Missing '{name_field_key.replace('_', ' ')}')"; print(f"Warning:{err_msg}"); generation_errors.append(err_msg); continue
            if current_missing_fields_keys:
                 err_msg = f"Manual {letter_type.capitalize()} Entry {i+1} ('{recipient_name}'): Missing fields: {', '.join(current_missing_fields_keys)}"; print(f"Warning:{err_msg}"); generation_errors.append(err_msg); continue

            name_suffix = recipient_name.replace(" ", "_").replace("/", "-").replace("\\", "-"); out_ts = datetime.datetime.now().strftime('%H%M%S%f')
            filename_suffix = secure_filename(f"manual_{letter_type}_{i+1}_{name_suffix}_{out_ts}")
            print(f"--- Processing Manual {letter_type.capitalize()} ({i+1}/{num_entries}): {recipient_name} ---")
            gen_success, docx_path, pdf_path, pdf_err_msg = generate_relieving_letter_web(actual_template_to_use, current_data_for_template, filename_suffix)

            added_files_for_record = False
            if gen_success:
                 if output_format in ['docx', 'both'] and docx_path and os.path.exists(docx_path): generated_files_paths.append(docx_path); added_files_for_record = True
                 if output_format in ['pdf', 'both']:
                     if pdf_path and os.path.exists(pdf_path):
                         generated_files_paths.append(pdf_path); added_files_for_record = True
                     else:
                         pdf_failures += 1
                         if pdf_err_msg: print(f"DEBUG: PDF failure reason for '{recipient_name}': {pdf_err_msg}")
                 if added_files_for_record:
                     successful_records_count += 1
                 else:
                     err = f"'{recipient_name}' (Manual {letter_type.capitalize()}): DOCX generated but no requested output format found/saved."
                     generation_errors.append(err); print(f"Warning: {err}")
            else:
                fail_reason = pdf_err_msg or 'Generation core function returned False.'
                generation_errors.append(f"'{recipient_name}' (Manual {letter_type.capitalize()}): Generation failed ({fail_reason})")

        print(f"DEBUG: Manual bulk {letter_type} finished. Success: {successful_records_count}/{num_entries}")
        if successful_records_count == 0:
             flash(f"No {letter_type} letters generated successfully from manual entries.", "danger");
             if generation_errors: flash("Errors: " + "; ".join(generation_errors), "warning");
             raise ValueError("No successful generation")

        zip_buffer = io.BytesIO(); print(f"DEBUG: Creating Manual {letter_type.capitalize()} ZIP ({output_format})...");
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zipf:
             for file_path in generated_files_paths:
                 if os.path.exists(file_path): zipf.write(file_path, os.path.basename(file_path))
                 else: print(f"Warning: File path listed for zipping but not found: {file_path}")

        zip_buffer.seek(0); success_msg = f"Generated documents for {successful_records_count}/{num_entries} manual {letter_type} entries (Format: {output_format.upper()})."; flash(success_msg, 'success')
        if generation_errors: flash("Issues: " + "; ".join(generation_errors), 'warning')
        if pdf_failures > 0 and output_format != 'docx': flash(f"Note: PDF failed/skipped for {pdf_failures} entries.", 'info')
        zip_filename_ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        zip_filename = f"Generated_Manual_Relieving_{output_format.upper()}_{zip_filename_ts}.zip"
        return send_file( zip_buffer, mimetype='application/zip', as_attachment=True, download_name=zip_filename )

    except ValueError:
        pass
    except Exception as route_err:
        print(f"CRITICAL Error /generate_bulk_manual_{letter_type}: {route_err}\n{traceback.format_exc()}"); flash("Unexpected server error during manual bulk generation.", "danger")
    finally:
        safe_cleanup(uploaded_temp_template_path_for_cleanup)
        if 'generated_files_paths' in locals() and zip_buffer: # Only cleanup if zip was created
             for p_cleanup in generated_files_paths: safe_cleanup(p_cleanup)

        if 'zip_buffer' not in locals() or not zip_buffer:
             return redirect(url_for('index', active_tab=active_tab_anchor))


# General Bulk Route (CSV/JSON for Offer/Relieving)
@app.route('/generate_bulk', methods=['POST'])
def handle_generate_bulk():
    uploaded_temp_template_path_for_cleanup = None
    data_path = None
    zip_buf = None
    letter_type = request.form.get('letter_type')
    source_type = request.form.get('source_type')
    output_format = request.form.get('output_format', 'both')
    # Determine active tab before any potential errors
    active_tab_anchor = f'{letter_type}-{source_type}-content' if letter_type and source_type else 'index'


    try:
        if not letter_type or letter_type not in ['offer', 'relieving']: flash('Invalid letter type selected.', 'danger'); raise ValueError("Invalid letter type")
        if not source_type or source_type not in ['csv', 'json']: flash('Invalid data source type selected.', 'danger'); raise ValueError("Invalid source type")
        if output_format not in ['both', 'docx', 'pdf']: flash('Invalid output format selected.', 'danger'); raise ValueError("Invalid output format")

        actual_template_to_use, uploaded_temp_template_path_for_cleanup = process_template_input(f"bulk_{letter_type}_{source_type}")
        if not actual_template_to_use: # process_template_input flashes its own errors
            raise ValueError("Template processing failed") # Will be caught and redirected correctly

        if 'data_file' not in request.files or not request.files['data_file'].filename:
            flash(f'No data file provided for bulk {letter_type} ({source_type}).', 'danger')
            raise ValueError("Missing data file")
        
        data_file = request.files['data_file']
        if not allowed_file(data_file.filename, ALLOWED_EXTENSIONS_DATA):
            flash(f'Invalid data file type for {source_type}.', 'danger')
            raise ValueError("Invalid data file type")
        
        file_ext = data_file.filename.rsplit('.',1)[1].lower()
        if file_ext != source_type:
            flash(f'Data file extension ".{file_ext}" does not match selected source type ".{source_type}".', 'danger')
            raise ValueError("Data file type mismatch")

        # --- Determine Letter-Specific Settings (CSV/JSON headers vs placeholders) ---
        # For CSV/JSON, the data file headers/keys become the *template placeholders* directly.
        # The `name_key_in_data` is the specific header/key in the CSV/JSON that holds the recipient's name.
        # The `name_ph_for_filename` is the *placeholder form* of that name key, used for messages/filenames if needed.
        if letter_type == 'offer': 
            name_key_in_data = 'candidate name' # Expected header in CSV/JSON for recipient's name
            name_ph_for_filename = '{candidate name}' 
            file_prefix = "Offer_Letter"
            gen_func = generate_offer_letter_web
        else: # relieving
            name_key_in_data = 'name' # Expected header in CSV/JSON for recipient's name
            name_ph_for_filename = '{name}'
            file_prefix = "Relieving_Letter"
            gen_func = generate_relieving_letter_web
        
        timestamp = datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')
        safe_data_filename = secure_filename(data_file.filename)
        temp_data_filename = f"upload_bulk_data_{letter_type}_{source_type}_{timestamp}_{safe_data_filename}"
        data_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_data_filename)
        data_file.save(data_path)
        print(f"DEBUG: Saved temporary data file: {data_path}")

        records = []; print(f"DEBUG: Processing {source_type.upper()} data for {letter_type.upper()}...");
        try:
            if source_type == 'csv':
                 df = pd.read_csv(data_path, dtype=str, keep_default_na=False)
                 # Normalize CSV headers: lowercase and strip spaces for robust matching
                 df.columns = [col.strip().lower() for col in df.columns]
                 # Check if the primary name key (normalized) exists in CSV headers
                 if name_key_in_data.lower().strip() not in df.columns:
                     raise ValueError(f"Missing required column header: '{name_key_in_data}' in CSV.")
                 
                 raw_records = df.to_dict('records')
                 for row_idx, row_data_dict in enumerate(raw_records):
                     # Create placeholder dict: {original_case_placeholder_from_csv_header}: value
                     # We need to map the original CSV headers to placeholders.
                     # This requires storing original headers before lowercasing, or careful reconstruction.
                     # For now, assume template placeholders match CSV headers *exactly* (case-insensitively for matching, but template needs exact)
                     # This implies template placeholders should be like {column name from csv}
                     
                     # Corrected approach: CSV headers are directly used as keys for template replacement (wrapped in {})
                     # Pandas df.to_dict('records') uses the (potentially modified) df.columns as keys.
                     # We need to use the *original* headers from the CSV if they differ in case/spacing from the template.
                     # The current df.columns are already normalized (lowercased, stripped).
                     # To simplify, we'll assume template placeholders are {lowercase_stripped_header_name}.
                     
                     # Re-read CSV to get original headers for placeholder creation
                     original_df = pd.read_csv(data_path, dtype=str, keep_default_na=False, nrows=0) # Read only headers
                     original_headers = [col.strip() for col in original_df.columns]

                     # Re-read full CSV with normalized headers for data access
                     df_data = pd.read_csv(data_path, dtype=str, keep_default_na=False)
                     df_data.columns = [col.strip().lower() for col in df_data.columns] # For access
                     
                     current_record_data_list = df_data.to_dict('records')
                     current_row_data_dict = current_record_data_list[row_idx]

                     # Build template data using original headers as placeholder keys
                     template_data = {}
                     for original_header in original_headers:
                         normalized_header = original_header.lower().strip()
                         placeholder_key = f"{{{original_header}}}" # Placeholder is {Original CSV Header}
                         template_data[placeholder_key] = str(current_row_data_dict.get(normalized_header, '')).strip()
                     records.append(template_data)


            elif source_type == 'json':
                 with open(data_path, 'r', encoding='utf-8') as f: raw_records_list = json.load(f)
                 if not isinstance(raw_records_list, list): raise ValueError("JSON data must be a list of objects.")
                 for obj_idx, raw_obj_dict in enumerate(raw_records_list):
                      if not isinstance(raw_obj_dict, dict): raise ValueError(f"JSON list item at index {obj_idx} must be an object.")
                      
                      # Check for name key (case-insensitive for robustness)
                      found_name_key_original_case = None
                      for k_obj in raw_obj_dict.keys():
                          if str(k_obj).strip().lower() == name_key_in_data.lower().strip():
                              found_name_key_original_case = k_obj; break
                      if not found_name_key_original_case:
                          raise ValueError(f"JSON object at index {obj_idx} is missing the required key: '{name_key_in_data}'.")

                      # JSON keys are used directly as placeholder keys (wrapped in {})
                      template_data = {f"{{{str(k).strip()}}}": str(v).strip() for k, v in raw_obj_dict.items()}
                      records.append(template_data)

            if not records: flash("No valid records found in data file.", "warning"); raise ValueError("No records")

        except Exception as parse_err: flash(f"Error parsing {source_type.upper()} data: {parse_err}", 'danger'); raise

        # --- Generation Loop ---
        generated_files_paths = []; gen_errors = []; pdf_fails = 0; success_count = 0; total_records = len(records)
        if output_format in ['pdf', 'both'] and not DOCX2PDF_AVAILABLE: flash("PDF conversion unavailable.", "warning")
        
        for i, record_data_for_template in enumerate(records):
            # Get recipient name using the placeholder form derived from `name_key_in_data`
            # For CSV, name_ph_for_filename was set e.g. {candidate name}. The actual placeholder in record_data_for_template
            # will be {Original CSV Header for candidate name}. We need to find this.
            # For JSON, name_ph_for_filename was set e.g. {name}. The actual placeholder in record_data_for_template
            # will be {Original JSON key for name}.

            recipient_name_from_record = ''
            # Try to find the name value by iterating through keys, comparing lowercased version without braces
            normalized_name_key_in_data = name_key_in_data.lower().strip()
            for ph_key, ph_value in record_data_for_template.items():
                # ph_key is like "{Some Header Name}"
                key_inside_braces = ph_key.strip()[1:-1] # Remove {}
                if key_inside_braces.lower().strip() == normalized_name_key_in_data:
                    recipient_name_from_record = ph_value.strip()
                    break
            
            if not recipient_name_from_record: 
                gen_errors.append(f"Record {i+1}: Missing or empty value for the name key ('{name_key_in_data}') in the {source_type.upper()} data."); continue;

            suffix_name = recipient_name_from_record.replace(" ","_").replace("/","-").replace("\\","-"); ts_suffix = datetime.datetime.now().strftime('%H%M%S%f')
            fname_suffix = secure_filename(f"bulk_{letter_type}_{source_type}_{i+1}_{suffix_name}_{ts_suffix}")
            print(f"--- Processing Bulk {letter_type.upper()} ({source_type.upper()}) ({i+1}/{total_records}): {recipient_name_from_record} ---")
            success, docx_p, pdf_p, pdf_err = gen_func(actual_template_to_use, record_data_for_template, fname_suffix)

            added = False
            if success:
                if output_format in ['docx','both'] and docx_p and os.path.exists(docx_p): generated_files_paths.append(docx_p); added = True
                if output_format in ['pdf','both']:
                    if pdf_p and os.path.exists(pdf_p): generated_files_paths.append(pdf_p); added = True
                    else: pdf_fails += 1;
                if added: success_count += 1
                else: gen_errors.append(f"'{recipient_name_from_record}': DOCX generated but no requested output format found/saved.")
            else: gen_errors.append(f"'{recipient_name_from_record}': Failed ({pdf_err or 'Unknown generation error'})")

        print(f"DEBUG: Bulk {letter_type} ({source_type}) finished. Success: {success_count}/{total_records}");
        if success_count == 0:
            flash(f"No {letter_type} documents generated successfully from {source_type} data.", "danger");
            if gen_errors: flash("Errors: "+"; ".join(gen_errors), "warning");
            raise ValueError("No successful generation")

        zip_buf = io.BytesIO(); print(f"DEBUG: Creating ZIP ({output_format}) for {letter_type} from {source_type}...")
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
             for fp in generated_files_paths:
                if os.path.exists(fp): zf.write(fp, os.path.basename(fp))
                else: print(f"Warning: File listed for zipping but not found: {fp}")
        zip_buf.seek(0); flash(f"Generated {success_count}/{total_records} {letter_type} documents from {source_type} (Format: {output_format.upper()}).", 'success')
        if gen_errors: flash("Issues: "+"; ".join(gen_errors), 'warning')
        if pdf_fails > 0 and output_format != 'docx': flash(f"Note: PDF failed/skipped for {pdf_fails} record(s).", 'info')
        zip_ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S"); zip_fname = f"Generated_{file_prefix}s_{source_type.upper()}_{output_format.upper()}_{zip_ts}.zip"
        return send_file(zip_buf, mimetype='application/zip', as_attachment=True, download_name=zip_fname)

    except ValueError: 
        pass 
    except Exception as route_err:
        print(f"CRITICAL Error /generate_bulk ({letter_type}/{source_type}): {route_err}\n{traceback.format_exc()}")
        flash(f"Unexpected server error during bulk {letter_type} generation from {source_type}.", "danger")
    finally:
        safe_cleanup(uploaded_temp_template_path_for_cleanup)
        safe_cleanup(data_path)
        if 'generated_files_paths' in locals() and zip_buf: # Only cleanup if zip was created
             for p_cleanup in generated_files_paths: safe_cleanup(p_cleanup)
        
        if 'zip_buf' not in locals() or not zip_buf: 
             return redirect(url_for('index', active_tab=active_tab_anchor))


# --- DOCX to PDF Converter Route ---
@app.route('/convert_to_pdf', methods=['POST'])
def handle_convert_to_pdf():
    active_tab_anchor = 'converter-content'
    if not DOCX2PDF_AVAILABLE: flash("PDF conversion capability unavailable on server.", "danger"); return redirect(url_for('index', active_tab=active_tab_anchor))
    if 'docx_file' not in request.files: flash('No DOCX file provided.', 'danger'); return redirect(url_for('index', active_tab=active_tab_anchor))
    file = request.files['docx_file']; original_filename = file.filename
    if original_filename == '': flash('No file selected.', 'danger'); return redirect(url_for('index', active_tab=active_tab_anchor))
    if not allowed_file(original_filename, ALLOWED_EXTENSIONS_DOCX): flash('Invalid file type (.docx required).', 'danger'); return redirect(url_for('index', active_tab=active_tab_anchor))
    
    temp_docx_path = None; generated_pdf_path = None
    try:
        safe_base = secure_filename(os.path.splitext(original_filename)[0]); ts = datetime.datetime.now().strftime('%Y%m%d%H%M%S%f')
        unique_suffix = f"{safe_base}_{ts}"
        temp_fn = f"upload_convert_in_{unique_suffix}.docx"; temp_docx_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_fn)
        gen_fn = f"converted_{unique_suffix}.pdf"; generated_pdf_path = os.path.join(app.config['GENERATED_FOLDER'], gen_fn)

        file.save(temp_docx_path); print(f"DEBUG: Converting {original_filename}...")
        if not docx_to_pdf_convert: raise RuntimeError("PDF Convert function unavailable (internal import issue).")
        docx_to_pdf_convert(temp_docx_path, generated_pdf_path) 
        if not os.path.exists(generated_pdf_path): raise RuntimeError("Conversion process completed, but the output PDF file was not found.")

        print(f"✅ Conversion successful: {gen_fn}"); flash(f"Successfully converted '{original_filename}' to PDF.", 'success')
        return redirect(url_for('index', 
                                converter_pdf_download=gen_fn, 
                                original_docx_filename=original_filename, 
                                active_tab=active_tab_anchor))
    except Exception as e:
        error_msg = f"Conversion failed for '{original_filename}': {str(e)}"; print(f"❌ Error in /convert_to_pdf: {e}\n{traceback.format_exc()}"); flash(error_msg, 'danger'); 
        safe_cleanup(generated_pdf_path); 
        return redirect(url_for('index', active_tab=active_tab_anchor))
    finally: 
        safe_cleanup(temp_docx_path)


# --- Download Route ---
@app.route('/download/<path:filename>')
def download_file(filename):
    safe_basename = secure_filename(filename)
    if not safe_basename or len(safe_basename) < 3:
        flash("Invalid download filename requested.", "danger")
        return redirect(url_for('index'))

    generated_dir = app.config['GENERATED_FOLDER']
    print(f"DEBUG: Download request for '{safe_basename}' from '{generated_dir}'")
    
    target_path = os.path.join(generated_dir, safe_basename)
    # Path traversal check
    if not os.path.abspath(target_path).startswith(os.path.abspath(generated_dir)):
        flash("Attempted to access file outside designated area.", "danger")
        print(f"SECURITY WARNING: Path traversal attempt for filename: {filename}")
        return redirect(url_for('index'))

    try:
        return send_from_directory(
            generated_dir,
            safe_basename,
            as_attachment=True
        )
    except FileNotFoundError:
         print(f"Error: Download file not found: {target_path}")
         flash(f"Sorry, the requested file '{safe_basename}' could not be found on the server.", "warning")
         return redirect(url_for('index'))
    except Exception as e:
         print(f"Error during download processing for '{safe_basename}': {e}")
         flash("An error occurred while preparing the file for download.", "danger")
         return redirect(url_for('index'))


# === Run the App ===
if __name__ == "__main__":
    print("--- Starting BKM Document Tools ---")
    is_debug_mode = os.environ.get("FLASK_DEBUG", "1") == "1" # Default to debug if not set
    app.debug = is_debug_mode

    print(f" * Environment: {'development' if app.debug else 'production'}")
    print(f" * Debug mode: {app.debug}")
    print(f" * Upload folder: {app.config['UPLOAD_FOLDER']}")
    print(f" * Generated files folder: {app.config['GENERATED_FOLDER']}")
    print(f" * User templates folder: {app.config['TEMPLATES_FOLDER']}")
    print(f" * PDF Conversion Available: {DOCX2PDF_AVAILABLE}")
    host_addr = os.environ.get("FLASK_RUN_HOST", '0.0.0.0')
    port_num = int(os.environ.get("FLASK_RUN_PORT", 5000))
    print(f" * Running on http://{host_addr}:{port_num} (Press CTRL+C to quit)")
    
    if app.debug:
        app.run(host=host_addr, port=port_num)
    else:
        # Example for Waitress (production)
        # from waitress import serve
        # serve(app, host=host_addr, port=port_num)
        # For simplicity in this context, still using app.run for non-debug,
        # but a proper WSGI server is recommended for production.
        app.run(host=host_addr, port=port_num)